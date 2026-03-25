import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

BASE = Path(__file__).resolve().parent
RESULTS_PATH = BASE / "assignment_results.json"
KMEANS_TXT = BASE / "7_kmeans_problem.txt"
PCA_TXT = BASE / "8_pca_problem.txt"
CODE_PATH = BASE / "solve_all_assignments.py"
OUT_PATH = BASE / "CA-I_Final_Submission_Report.docx"


def set_default_font(document):
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(12)


def add_title_page(doc):
    p = doc.add_paragraph("MSPMS")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True

    p = doc.add_paragraph("DEOGIRI INSTITUTE OF ENGINEERING & MANAGEMENT STUDIES, AURANGABAD")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True

    p = doc.add_paragraph("DEPARTMENT OF CSE (AI & ML)")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph("A.Y. 2025-26, SEM-II")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")
    p = doc.add_paragraph("CONTINUOUS ASSESSMENT-I (CA-I) SUBMISSION REPORT")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True

    p = doc.add_paragraph("SUBJECT: ADVANCED MACHINE LEARNING (BTAIC602)")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")
    doc.add_paragraph("Student Name: __________________________")
    doc.add_paragraph("Roll No.: ______________________________")
    doc.add_paragraph("Class: T.Y. CSE (AI & ML)")
    doc.add_paragraph("Date of Submission: 25/03/2026")
    doc.add_page_break()


def add_problem_statements(doc, kmeans_text, pca_text):
    doc.add_heading("1. Problem Statements", level=1)

    doc.add_heading("1.1 K-Means Clustering Problem Statements", level=2)
    for line in kmeans_text.splitlines():
        line = line.strip()
        if line:
            doc.add_paragraph(line)

    doc.add_heading("1.2 PCA Problem Statements", level=2)
    for line in pca_text.splitlines():
        line = line.strip()
        if line:
            doc.add_paragraph(line)


def add_kmeans_solution(doc, kmeans):
    doc.add_heading("2. K-Means Clustering Solutions", level=1)

    common = doc.add_paragraph()
    common.add_run("Common Methodology: ").bold = True
    common.add_run(
        "Data cleaning, handling missing values, categorical encoding for mixed data, feature scaling, "
        "K selection by scree/silhouette, validation with hierarchical clustering and ARI."
    )

    mapping = [
        ("1_airlines", "Problem 1: Airlines Dataset (EastWestAirlines.xlsx)"),
        ("2_crime", "Problem 2: Crime Dataset (crime_data.csv)"),
        ("3_insurance", "Problem 3: Insurance Dataset (Insurance Dataset.csv)"),
        ("4_telco", "Problem 4: Telecom Mixed Dataset (Telco_customer_churn.xlsx)"),
        ("5_autoinsurance", "Problem 5: AutoInsurance Mixed Dataset (Autoinsurance.csv)"),
    ]

    for key, title in mapping:
        v = kmeans[key]
        doc.add_heading(title, level=2)
        doc.add_paragraph(f"Dataset shape: {v['shape'][0]} rows x {v['shape'][1]} columns")
        doc.add_paragraph(f"Total missing values: {v['missing_values']}")
        doc.add_paragraph(f"Optimum number of clusters (K): {v['scree']['best_k']}")
        doc.add_paragraph(f"KMeans silhouette score: {v['kmeans_silhouette']}")
        doc.add_paragraph(f"Hierarchical silhouette score: {v['hierarchical_silhouette']}")
        doc.add_paragraph(f"ARI (KMeans vs Hierarchical): {v['ari_kmeans_vs_hierarchical']}")
        doc.add_paragraph(f"KMeans cluster sizes: {v['kmeans_cluster_sizes']}")

        p = doc.add_paragraph()
        p.add_run("Inference: ").bold = True
        if key == "1_airlines":
            p.add_run("Multiple customer segments are formed, including high-value flyer groups and low-engagement travelers.")
        elif key == "2_crime":
            p.add_run("Two broad state groups are identified: relatively lower-crime and higher-crime profiles.")
        elif key == "3_insurance":
            p.add_run("Customers split into standard and high-value/high-claim segments useful for policy targeting.")
        elif key == "4_telco":
            p.add_run("Two major segments emerge based on tenure and revenue; long-tenure/high-revenue users are clearly separated.")
        else:
            p.add_run("Segments differ by income and claim burden; lower silhouette indicates additional feature engineering could improve separability.")


def add_pca_solution(doc, pca):
    doc.add_heading("3. PCA Problem Statement Solution (Heart Disease)", level=1)

    doc.add_paragraph(f"Dataset shape: {pca['shape'][0]} rows x {pca['shape'][1]} columns")
    doc.add_paragraph(f"Missing values: {pca['missing_values']}")
    doc.add_paragraph(f"Target distribution: {pca['target_distribution']}")
    doc.add_paragraph(f"Optimum K from scree/silhouette: {pca['scree']['best_k']}")

    doc.add_paragraph("PCA analysis:")
    doc.add_paragraph(f"- Explained variance ratio (PC1, PC2, PC3): {pca['pca_explained_variance_ratio']}")
    doc.add_paragraph(f"- Cumulative variance by first 3 PCs: {pca['pca_cumulative_variance_3']}")
    doc.add_paragraph("- New 3-PC dataset created: heart_pca_3_components.csv")

    doc.add_paragraph("Clustering quality comparison:")
    doc.add_paragraph(f"- Original KMeans silhouette: {pca['original_kmeans_silhouette']}")
    doc.add_paragraph(f"- Original Hierarchical silhouette: {pca['original_hierarchical_silhouette']}")
    doc.add_paragraph(f"- PCA KMeans silhouette: {pca['pca_kmeans_silhouette']}")
    doc.add_paragraph(f"- PCA Hierarchical silhouette: {pca['pca_hierarchical_silhouette']}")
    doc.add_paragraph(f"- ARI (KMeans original vs PCA): {pca['ari_kmeans_original_vs_pca']}")
    doc.add_paragraph(f"- ARI (Hierarchical original vs PCA): {pca['ari_hierarchical_original_vs_pca']}")

    p = doc.add_paragraph()
    p.add_run("Inference: ").bold = True
    p.add_run("KMeans clusters are highly similar before and after PCA, and silhouette scores improve on PCA-transformed space.")


def add_business_impact(doc):
    doc.add_heading("4. Business Impact", level=1)
    points = [
        "Customer/patient segmentation supports targeted decision-making.",
        "Clustering reveals hidden patterns across behavior and risk features.",
        "PCA reduces dimensionality and improves interpretability with lower computational complexity.",
        "The solution helps prioritize interventions, campaigns, and analytical monitoring.",
    ]
    for pt in points:
        doc.add_paragraph(pt, style="List Bullet")


def add_code_section(doc, code_text):
    doc.add_heading("5. Python Code Used", level=1)
    doc.add_paragraph("The following code was used to compute all assignment results:")

    # Put code in fixed-width style by setting run font.
    code_para = doc.add_paragraph()
    run = code_para.add_run(code_text)
    run.font.name = "Courier New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
    run.font.size = Pt(9)


def add_conclusion(doc):
    doc.add_heading("6. Conclusion", level=1)
    doc.add_paragraph(
        "All required problem statements from K-Means Clustering and PCA have been solved with proper preprocessing, "
        "model selection, validation, and interpretation. The report includes both problem statements and final solutions "
        "in submission-ready format."
    )


def main():
    with open(RESULTS_PATH, "r", encoding="utf-8") as f:
        results = json.load(f)

    kmeans_text = KMEANS_TXT.read_text(encoding="utf-8")
    pca_text = PCA_TXT.read_text(encoding="utf-8")
    code_text = CODE_PATH.read_text(encoding="utf-8")

    doc = Document()
    set_default_font(doc)

    add_title_page(doc)
    add_problem_statements(doc, kmeans_text, pca_text)
    add_kmeans_solution(doc, results["kmeans_pdf_problem_statements"])
    add_pca_solution(doc, results["pca_pdf_problem_statement"]["heart_disease"])
    add_business_impact(doc)
    add_code_section(doc, code_text)
    add_conclusion(doc)

    doc.save(str(OUT_PATH))
    print(f"Created report: {OUT_PATH}")


if __name__ == "__main__":
    main()
