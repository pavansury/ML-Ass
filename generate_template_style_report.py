import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

BASE = Path(__file__).resolve().parent
RESULTS_PATH = BASE / "assignment_results.json"
OUT_PATH = BASE / "CA-I_Final_Submission_Report.docx"
CODE_PATH = BASE / "solve_all_assignments.py"


def set_font(doc):
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(12)


def add_heading_center(doc, text, bold=True):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if p.runs:
        p.runs[0].bold = bold


def add_bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


def main():
    with open(RESULTS_PATH, "r", encoding="utf-8") as f:
        data = json.load(f)

    k = data["kmeans_pdf_problem_statements"]
    p = data["pca_pdf_problem_statement"]["heart_disease"]

    doc = Document()
    set_font(doc)

    # Template-like header
    add_heading_center(doc, "DATA ANALYSIS REPORT", bold=True)
    doc.add_paragraph("Name: Pavan Sajjan Suryawanshi")
    doc.add_paragraph("Roll No: AI3078")
    doc.add_paragraph("Topic: K-MEANS CLUSTERING AND PCA")
    add_heading_center(doc, "ASSIGNMENT", bold=True)

    doc.add_paragraph("TYPE OF EXAM: APPLY CLUSTERING ALGORITHMS AND DIMENSION REDUCTION")
    doc.add_paragraph("ANALYZE DATA FOR ALGORITHMS")
    doc.add_paragraph("1. K-MEANS CLUSTERING")
    doc.add_paragraph("2. HIERARCHICAL CLUSTERING")
    doc.add_paragraph("3. PRINCIPAL COMPONENT ANALYSIS (PCA)")

    doc.add_paragraph("Problem Statement")
    doc.add_paragraph(
        "The aim of this assignment is to analyze multiple business datasets using unsupervised learning techniques. "
        "For K-means tasks, the objective is to identify optimum clusters and derive business insights. "
        "For PCA task, the objective is to reduce dimensionality and compare clustering performance before and after PCA."
    )

    # Template-like numbered sections
    doc.add_paragraph("1. Import Libraries")
    doc.add_paragraph("2. Load Dataset")
    doc.add_paragraph("3. Data Preprocessing")
    doc.add_paragraph("4. Feature Scaling / Encoding")
    doc.add_paragraph("5. K-MEANS CLUSTERING")
    doc.add_paragraph("6. HIERARCHICAL CLUSTERING")
    doc.add_paragraph("7. PCA (for Heart Disease dataset)")
    doc.add_paragraph("8. Evaluation and Insights")

    doc.add_heading("DATA ANALYSIS", level=1)

    doc.add_heading("1. Data Collection & Cleaning", level=2)
    add_bullet(doc, "Used datasets available in assignment folder for Airlines, Crime, Insurance, Telco, AutoInsurance, and Heart Disease.")
    add_bullet(doc, "Handled missing values using imputation for mixed datasets (especially telecom).")
    add_bullet(doc, "Removed ID-like columns where required (e.g., Customer ID, ID#, Customer).")
    add_bullet(doc, "Converted categorical variables using one-hot encoding for mixed-data clustering.")
    add_bullet(doc, "Standardized features before clustering and PCA.")

    doc.add_heading("2. Data Analysis", level=2)
    doc.add_paragraph("I performed exploratory analysis and clustering validation using silhouette score and ARI.")
    add_bullet(doc, f"Airlines optimum K = {k['1_airlines']['scree']['best_k']} with silhouette = {k['1_airlines']['kmeans_silhouette']}.")
    add_bullet(doc, f"Crime optimum K = {k['2_crime']['scree']['best_k']} with silhouette = {k['2_crime']['kmeans_silhouette']}.")
    add_bullet(doc, f"Insurance optimum K = {k['3_insurance']['scree']['best_k']} with silhouette = {k['3_insurance']['kmeans_silhouette']}.")
    add_bullet(doc, f"Telco optimum K = {k['4_telco']['scree']['best_k']} with silhouette = {k['4_telco']['kmeans_silhouette']}.")
    add_bullet(doc, f"AutoInsurance optimum K = {k['5_autoinsurance']['scree']['best_k']} with silhouette = {k['5_autoinsurance']['kmeans_silhouette']}.")

    doc.add_heading("3. Splitting and Transformation Logic", level=2)
    doc.add_paragraph(
        "Since this is unsupervised learning, no train-test split is mandatory for clustering quality measurement. "
        "Instead, full-data clustering with internal validation (silhouette) and cross-method comparison (ARI) was used."
    )

    doc.add_heading("4. Algorithms Used", level=2)
    doc.add_paragraph("A. K-Means Clustering")
    add_bullet(doc, "Used scree/elbow trend and silhouette score to pick optimum K.")
    add_bullet(doc, "Generated cluster sizes and profile means for interpretation.")

    doc.add_paragraph("B. Hierarchical Clustering")
    add_bullet(doc, "Applied Ward linkage with same K for comparison.")
    add_bullet(doc, "Compared agreement with K-means using Adjusted Rand Index.")

    doc.add_paragraph("C. Principal Component Analysis (Heart Disease)")
    add_bullet(doc, f"Best K before PCA = {p['scree']['best_k']}.")
    add_bullet(doc, f"Explained variance ratio (PC1, PC2, PC3) = {p['pca_explained_variance_ratio']}.")
    add_bullet(doc, f"Cumulative variance by first 3 PCs = {p['pca_cumulative_variance_3']}.")
    add_bullet(doc, f"ARI K-means (before vs after PCA) = {p['ari_kmeans_original_vs_pca']} (high similarity).")

    doc.add_heading("5. Evaluation Metrics", level=2)
    add_bullet(doc, "Silhouette Score for cluster compactness and separation.")
    add_bullet(doc, "Adjusted Rand Index (ARI) to compare clustering consistency across methods and PCA states.")

    doc.add_heading("6. Problem-Wise Inference", level=2)
    add_bullet(doc, "Airlines: identified distinct customer loyalty segments including high-value flyers.")
    add_bullet(doc, "Crime: states split into lower-crime and higher-crime profiles.")
    add_bullet(doc, "Insurance: strong separation between standard and high-value/high-claim customers.")
    add_bullet(doc, "Telco: two major groups differ significantly by tenure and revenue, useful for churn-focused action.")
    add_bullet(doc, "AutoInsurance: segments differ by claim burden and income; further feature engineering can improve separability.")
    add_bullet(doc, "PCA: dimensionality reduction preserved K-means structure while improving silhouette after transformation.")

    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph(
        "All problem statements were solved as per assignment instructions. The clustering solutions and PCA comparison "
        "provide actionable insights for segmentation and decision support. For the heart disease case, PCA retained core "
        "structure while improving cluster quality in reduced space."
    )

    doc.add_heading("APPENDIX: Python Code", level=1)
    code_text = CODE_PATH.read_text(encoding="utf-8")
    code_par = doc.add_paragraph()
    run = code_par.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)

    doc.save(str(OUT_PATH))
    print(f"Updated report in template format: {OUT_PATH}")


if __name__ == "__main__":
    main()
